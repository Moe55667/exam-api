from django.shortcuts import render
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser,FormParser
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from django.core.files.uploadedfile import InMemoryUploadedFile
from io import BytesIO
from docx import Document
from .serializers import GenerateExamSerializer,CorrectExamSerializer,VisionExamSerializer,StudentMarkSerializer,CustomUserSerializer
from rest_framework import generics, status
from .models import QuestionDetail, StudentReview
from django.http import HttpResponse
import csv
from django.contrib.auth.models import User
from rest_framework.permissions import IsAdminUser
from openai import OpenAI
import openai
import time
from .utils import extract_review_details,save_review_to_db,extract_exam_text
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
import os

from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize OpenAI client

client = OpenAI()
assistant1 = "asst_GCINmHDsrAbyiJ1sOq0JA8mX"
assistant2 = "asst_nRK4FSX9WAbFWZlQdLAVVQyr"

# Create your views here.

class GenerateExamAPIView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = GenerateExamSerializer
    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data)
        
        if serializer.is_valid():
            uploaded_file = serializer.validated_data['file']
            exam_chapters = serializer.validated_data['exam_chapters']
            exam_parts = serializer.validated_data['exam_parts']
            each_part_question_numbers = serializer.validated_data['each_part_question_numbers']
            exam_difficulty = serializer.validated_data['exam_difficulty']
            exam_level = serializer.validated_data['exam_level']

            # Read the uploaded file as bytes
            uploaded_file.seek(0)
            file_bytes = uploaded_file.read()
            # Upload the user-provided file to OpenAI
            message_file = client.files.create(
                file=(uploaded_file.name, file_bytes), purpose="assistants"
            )


            # Create a thread
            thread = client.beta.threads.create(
                messages=[
                    {
                        "role": "user",
                        "content": f"Please generate an exam based on the book. Exam should contain the following chapters : {exam_chapters}, The exam should contain the following parts  which are {', '.join(exam_parts)},  each part will have {each_part_question_numbers}. The exam level is {exam_level} and the difficulty is {exam_difficulty}. Your output response should only be the exam text, without any extra instructions or answers. WARNING : NEVER IN ANY CASE INCLUDE ANSWERS",
                        "attachments": [
                            {"file_id": message_file.id, "tools": [{"type": "file_search"}]}
                        ],
                    }
                ]
            )

            # Start run
            run = client.beta.threads.runs.create_and_poll(
                thread_id=thread.id, assistant_id=assistant1
            )

            # Outputting message
            messages = list(client.beta.threads.messages.list(thread_id=thread.id, run_id=run.id))

            # check if messages is there 
            if messages:
                message_content = messages[0].content[0].text
                    # Generate DOCX file
                doc = Document()
                doc.add_heading('Generated Exam', level=1)
                doc.add_paragraph(message_content.value)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                

                # Create an in-memory file
                in_memory_file = InMemoryUploadedFile(
                    buffer, None, 'generated_exam.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', buffer.tell(), None
                )

                response = Response(
                    {'message': message_content},
                    status=status.HTTP_200_OK
                )
                response['Content-Disposition'] = f'attachment; filename=generated_exam.docx'
                response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                response.content = in_memory_file.read()
                return response
            else :
                response = Response(
                    {'message': "No message is found Out of Money"},
                    status=status.HTTP_200_OK
                )
                return response
           
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class CorrectExamAPIView(APIView):
    def post(self, request, *args, **kwargs):
        serializer = CorrectExamSerializer(data=request.data)
        if serializer.is_valid():
            original_file = serializer.validated_data['original_file']
            student_files = serializer.validated_data['student_files']

            # Ensure original file has a supported extension
            if not original_file.name.endswith(('.pdf', '.docx')):
                return Response({'error': f'Unsupported file extension for {original_file.name}'}, status=status.HTTP_400_BAD_REQUEST)

            vector_store = client.beta.vector_stores.create(name="Book vector store")

            # Read the original file content correctly and prepare for upload
            original_file_stream = original_file.read()

            # Use the upload and poll SDK helper to upload the files, add them to the vector store,
            # and poll the status of the file batch for completion.
            file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
                vector_store_id=vector_store.id,
                files=[
                    (original_file.name, original_file_stream),
                ]
            )

            # Extract text from the images
            imgs_list = []
            for img in student_files:
                img_stream = img.read()
                imgs_list.append(img_stream)

            student_exam_papers = extract_exam_text(imgs_list)
            print(student_exam_papers)

            # Prepare the prompt for the AI assistant
            prompt = f'''
                Please review this student exam {student_exam_papers} using the uploaded corrected exam file {original_file.name} as the reference material to identify any errors.
            Provide the answer in a structured format as specified below.

            Instructions:
            1. Review the student exam file using the uploaded file as a reference to detect errors.
            2. Determine the student's final score and specify the number of questions they answered incorrectly.
            3. Respond in a structured manner to improve clarity. Avoid manually evaluating each question individually.

            Response Format should be this json like format please :
            {{
            "review": {{
                "student_name": "[Extract Student Name from student exam text passed]",
                "book_name": "[Subject name]",
                "results": {{
                    "final_score": [Number of Correct Answers and Final Score Marks],
                    "incorrect_answers_count": [Number of Incorrect Answers],
                    "incorrect_answers_details": [
                        {{
                            "question_number": [Question Number ],
                            "reason": "[reason for the incorrect answer]",
                            "correct_answer": "[Correct Answer of the question]"
                        }}
                        // Repeat for each incorrect answer
                    ]
                }}
            }}
            }}

            WARNING: Make sure incorrect_answers_count and number of questions in incorrect_answers_details match for consistent reasons for example if incorrect_answers_count are 4 incorrect_answers_details should be 4 this is must make sure you NEVER FORGET

            '''
            
            thread = client.beta.threads.create(
                messages=[ { "role": "user", "content": prompt} ],
                tool_resources={
                    "file_search": {
                    "vector_store_ids": [vector_store.id]
                    }
                }
            )
            # Run the assistant
            run = client.beta.threads.runs.create_and_poll(
                thread_id=thread.id, assistant_id=assistant2
            )
            print('status',run.status)
            # Poll until the run status is completed
            while run.status != "completed":
                time.sleep(2)  # Adding a delay to avoid rapid polling
                run = client.beta.threads.runs.retrieve(run.id, thread_id=thread.id)

            messages = list(client.beta.threads.messages.list(thread_id=thread.id, run_id=run.id))
            if messages:
                message_content = messages[0].content[0].text
                annotations = message_content.annotations
                citations = []
                for index, annotation in enumerate(annotations):
                    message_content.value = message_content.value.replace(annotation.text, f"[{index}]")
                    if file_citation := getattr(annotation, "file_citation", None):
                        cited_file = client.files.retrieve(file_citation.file_id)
                        citations.append(f"[{index}] {cited_file.filename}")
                
                # Structured response
                structured_data = extract_review_details(message_content.value)
                save_review_to_db(structured_data, self.request.user)
                # Extract and structure the response
                # Return the status and file counts of the batch to see the result of this operation.
                response = Response({
                    'status': file_batch.status,
                    'file_counts': file_batch.file_counts,
                    'response': structured_data,
                }, status=status.HTTP_200_OK)
                # Add this line to set the CORS header
                return response
            else:
                return Response({
                    'status': file_batch.status,
                    'file_counts': file_batch.file_counts,
                    'response': "No response"
                }, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class RecordStudentMarkView(generics.ListCreateAPIView):
    serializer_class = StudentMarkSerializer

    def get_queryset(self):
        return StudentReview.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class DownloadStudentMarksCsvView(APIView):
    serializer_class = StudentMarkSerializer

    def get(self, request):
        marks = StudentReview.objects.filter(user=request.user)
        
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="student_marks.csv"'

        writer = csv.writer(response)
        writer.writerow(['Student Name', 'Book Name', 'Final Score', 'Incorrect Answers', 'Question Number', 'Status', 'Correct Version'])
        
        for mark in marks:
            detailed_reviews = QuestionDetail.objects.filter(student_review=mark)
            for detail in detailed_reviews:
                writer.writerow([
                    mark.student_name,
                    mark.book_name,
                    mark.final_score,
                    mark.incorrect_answers_count,
                    detail.question_number,
                    detail.status,
                    detail.correct_version
                ])

        return response
    
class StudentMarkDetailView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = StudentMarkSerializer

    def get_queryset(self):
        return StudentReview.objects.filter(user=self.request.user)
    
class UserListCreateView(generics.ListCreateAPIView):
    queryset = User.objects.all()
    serializer_class = CustomUserSerializer
    permission_classes = [IsAdminUser]

class UserRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    queryset = User.objects.all()
    serializer_class = CustomUserSerializer
    permission_classes = [IsAdminUser]












